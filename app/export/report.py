"""
Esportazione del colloquio in Word (.docx) e testo semplice (.txt).

Per il PDF ci si appoggia alla funzione "Salva come PDF" di Word o
LibreOffice: evita di trascinare nel pacchetto una libreria di
impaginazione aggiuntiva solo per questo.
"""
from __future__ import annotations

import logging
from datetime import datetime
from pathlib import Path

from app import config
from app.storage.db import Interview

log = logging.getLogger(__name__)


def _safe_filename(name: str) -> str:
    cleaned = "".join(c for c in name if c.isalnum() or c in "-_ ").strip()
    # Windows rifiuta i nomi troppo lunghi: un nome candidato incollato
    # per errore da un curriculum non deve far fallire l'esportazione.
    return (cleaned.replace(" ", "_") or "colloquio")[:60]


def _base_name(interview: Interview) -> str:
    # Usiamo la data del colloquio, non quella odierna: esportando dallo
    # storico un colloquio di mesi fa, il file deve portarne la data.
    try:
        when = datetime.fromisoformat(interview.created_at)
    except Exception:
        when = datetime.now()
    return f"{_safe_filename(interview.candidate_name)}_{when.strftime('%Y%m%d-%H%M%S')}"


def _unique_path(directory: Path, base: str, suffix: str) -> Path:
    """
    Evita di sovrascrivere un file gia' esistente senza avvisare.

    Esportando due volte lo stesso colloquio, il secondo file prende un
    numero progressivo invece di cancellare il primo.
    """
    path = directory / f"{base}{suffix}"
    counter = 2
    while path.exists():
        path = directory / f"{base}_{counter}{suffix}"
        counter += 1
    return path


def _format_duration(seconds: float) -> str:
    seconds = max(0, int(seconds or 0))
    minutes, secs = divmod(seconds, 60)
    return f"{minutes} min {secs:02d} s"


def _transcript_lines(interview: Interview, labels: dict[str, str]) -> list[str]:
    if interview.segments:
        return [
            f"{labels.get(seg.get('speaker', ''), seg.get('speaker', ''))}: {seg.get('text', '')}"
            for seg in interview.segments
        ]
    return (interview.transcript or "").splitlines()


def export_txt(
    interview: Interview, labels: dict[str, str], directory: Path | None = None
) -> Path:
    directory = directory or config.EXPORTS_DIR
    directory.mkdir(parents=True, exist_ok=True)
    path = _unique_path(directory, _base_name(interview), ".txt")

    lines = [
        f"COLLOQUIO - {interview.candidate_name}",
        f"Posizione: {interview.role or 'non indicata'}",
        f"Data: {interview.display_date}",
        f"Durata: {_format_duration(interview.duration_seconds)}",
        f"Lingua rilevata: {interview.detected_language or 'non rilevata'}",
        "",
        "=== REPORT ===",
        interview.report or "(nessun report)",
    ]
    if interview.notes:
        lines += ["", "=== NOTE ===", interview.notes]
    lines += ["", "=== TRASCRIZIONE ==="] + _transcript_lines(interview, labels)

    # La firma UTF-8 iniziale serve a Blocco note ed Excel per mostrare
    # correttamente le lettere accentate.
    path.write_text("\n".join(lines), encoding="utf-8-sig")
    log.info("Esportato: %s", path)
    return path


def export_docx(
    interview: Interview, labels: dict[str, str], directory: Path | None = None
) -> Path:
    from docx import Document
    from docx.shared import Pt, RGBColor

    directory = directory or config.EXPORTS_DIR
    directory.mkdir(parents=True, exist_ok=True)
    path = _unique_path(directory, _base_name(interview), ".docx")

    document = Document()
    document.add_heading(f"Colloquio - {interview.candidate_name}", level=0)

    meta = document.add_paragraph()
    meta.add_run("Posizione: ").bold = True
    meta.add_run(f"{interview.role or 'non indicata'}\n")
    meta.add_run("Data: ").bold = True
    meta.add_run(f"{interview.display_date}\n")
    meta.add_run("Durata: ").bold = True
    meta.add_run(f"{_format_duration(interview.duration_seconds)}\n")
    meta.add_run("Lingua rilevata: ").bold = True
    meta.add_run(interview.detected_language or "non rilevata")

    document.add_heading("Report", level=1)
    if not interview.used_llm and interview.report:
        note = document.add_paragraph(
            "Resoconto generato senza modello linguistico locale."
        )
        for run in note.runs:
            run.italic = True
            run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    for block in (interview.report or "").split("\n"):
        if block.strip():
            document.add_paragraph(block.strip())

    if interview.notes:
        document.add_heading("Note", level=1)
        for block in interview.notes.split("\n"):
            if block.strip():
                document.add_paragraph(block.strip(), style="List Bullet")

    document.add_heading("Trascrizione completa", level=1)
    for line in _transcript_lines(interview, labels):
        if not line.strip():
            continue
        paragraph = document.add_paragraph()
        if ":" not in line:
            # Riga senza interlocutore (trascrizione di una versione
            # precedente): va riportata cosi' com'e'.
            plain = paragraph.add_run(line.strip())
            plain.font.size = Pt(9)
            continue
        speaker, _, text = line.partition(":")
        run = paragraph.add_run(f"{speaker}: ")
        run.bold = True
        run.font.size = Pt(9)
        body = paragraph.add_run(text.strip())
        body.font.size = Pt(9)

    document.save(str(path))
    log.info("Esportato: %s", path)
    return path
